"""Extract text from DOCX resumes using python-docx."""
from __future__ import annotations

from pathlib import Path

from docx import Document


class DOCXParseError(Exception):
    pass


def parse_docx(path: str | Path) -> str:
    """Extract raw text from a DOCX file, including table cells."""
    path = Path(path)
    try:
        doc = Document(str(path))
        parts: list[str] = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)
    except Exception as exc:
        raise DOCXParseError(f"Failed to parse DOCX {path}: {exc}") from exc
